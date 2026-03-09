from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# === Page margins ===
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(2)
pf.line_spacing = Pt(12)

# ---------------------------------------------------------------------------
# HELPER FUNCTIONS
# ---------------------------------------------------------------------------
navy = RGBColor(0x1B, 0x3A, 0x5C)
dark = RGBColor(0x2D, 0x2D, 0x2D)
medium = RGBColor(0x4A, 0x4A, 0x4A)

def add_horizontal_line(doc, color='1B3A5C', weight=1.5):
    """Add a colored horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(weight * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(text):
    """Add a section heading with line underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = navy
    run.font.name = 'Calibri'
    add_horizontal_line(doc)

def role_header(title, dates):
    """Add a role title with dates right-aligned on same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    # Use a tab stop for right-alignment of dates
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.2), alignment=2)  # RIGHT aligned
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = navy
    run.font.name = 'Calibri'
    p.add_run('\t')
    run2 = p.add_run(dates)
    run2.font.size = Pt(10)
    run2.font.color.rgb = medium
    run2.font.name = 'Calibri'

def org_line(text):
    """Add organization / department line in italic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = medium
    run.font.name = 'Calibri'

def role_summary(text):
    """Add a brief role summary paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = dark
    run.font.name = 'Calibri'

def sub_heading(text):
    """Add a sub-heading within a role (e.g., 'AI-ENABLED STRATEGY')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = navy
    run.font.name = 'Calibri'

def bullet(text, indent=0.25):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = dark
    run.font.name = 'Calibri'

def simple_para(text, bold=False, size=10, color=dark, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p


# =========================================================================
# NAME HEADER
# =========================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('JAQUAN K. LEVONS')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = navy
run.font.name = 'Calibri'

# Contact line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('Titusville, NJ  \u2022  609.240.1241  \u2022  jaquan.levons@gmail.com  \u2022  linkedin.com/in/jaquan-levons')
run.font.size = Pt(10)
run.font.color.rgb = medium
run.font.name = 'Calibri'

add_horizontal_line(doc, weight=2)

# =========================================================================
# EXECUTIVE SUMMARY
# =========================================================================
section_heading('EXECUTIVE SUMMARY')

summary = (
    'Digital transformation and AI-enablement leader with 22+ years at Bristol Myers Squibb, '
    'progressing from hands-on drug product scientist through automation pioneer and digital '
    'strategist to Scientific Director overseeing enterprise AI/ML, modeling, and data transformation '
    'programs across Drug Product Development. Architect of BMS\u2019s three-tier automation value '
    'framework and the GPS Digital Moonshot 5-year vision \u2014 bridging IT technology experts and '
    'drug development scientists to deliver measurable portfolio value. Track record of translating '
    'ambiguous digital mandates into funded, governed programs with quantified impact: $3M+ '
    'automation platform serving 80+ assets, GenAI strategy yielding 40+ prioritized use cases, and '
    'predictive modeling initiatives saving ~2 FTE while accelerating decisions across 6 major '
    'programs. Deep pharmaceutical value chain expertise spanning preformulation through commercial '
    'tech transfer across small molecule, biologics, and gene therapy modalities.'
)
role_summary(summary)

# =========================================================================
# CORE STRENGTHS
# =========================================================================
section_heading('CORE STRENGTHS')

strengths = [
    ('Digital & AI Strategy for Drug Development', 'IT\u2013Science Bridge Building & Stakeholder Alignment'),
    ('AI/ML & GenAI Program Leadership', 'Portfolio Governance & ROI Frameworks'),
    ('Enterprise Data Strategy & Architecture', 'Cross-Functional Program Execution'),
    ('Change Management & Adoption at Scale', 'People Leadership & Talent Development'),
    ('Agile/Scrum Product Ownership', 'Pharmaceutical Value Chain (R&D \u2192 Manufacturing)'),
]

# Build as a two-column table with no borders
table = doc.add_table(rows=len(strengths), cols=2)
table.autofit = True
for i, (left, right) in enumerate(strengths):
    for j, text in enumerate([left, right]):
        cell = table.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run('\u2022  ' + text)
        run.font.size = Pt(10)
        run.font.color.rgb = dark
        run.font.name = 'Calibri'
    # Remove borders
    for j in range(2):
        tc = table.cell(i, j)._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

# =========================================================================
# PROFESSIONAL EXPERIENCE
# =========================================================================
section_heading('PROFESSIONAL EXPERIENCE')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('BRISTOL MYERS SQUIBB')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = dark
run.font.name = 'Calibri'
p.add_run('\t')
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(6.2), alignment=2)
run2 = p.add_run('2003\u2013Present')
run2.font.size = Pt(10)
run2.font.color.rgb = medium
run2.font.name = 'Calibri'

# --- ROLE 1: Scientific Director ---
role_header('Scientific Director, Automation, Digital & AI Transformation', 'Jun 2022\u2013Present')
org_line('Drug Product Development / PD Development Excellence (DevEx)')
role_summary(
    'Enterprise-level accountability for digital strategy, AI/ML adoption, predictive modeling, and '
    'data transformation across Drug Product Development \u2014 bridging IT and drug development '
    'scientists to deliver measurable portfolio value across biologics, small molecule, and gene therapy programs.'
)

sub_heading('AI-ENABLED STRATEGY & EXECUTION')
bullet('Developed and executed the DPD Digital, AI & Innovation Strategy encompassing predictive modeling, GenAI/LLM integration, structured data capture, and enterprise tool rationalization \u2014 directly aligned with PD DevEx and GPS strategic priorities.')
bullet('Led the GenAI/LLM strategy for DPD: collected and assessed 40+ use cases against a prioritization framework, piloted DocGenAI (regulatory document drafting) and ECL coding bot, implemented office hours and biweekly newsletter for AI literacy, and proposed a strategic roadmap of 5 quick wins and 4 big bets.')
bullet('Co-led the Bio In Silico DevEx Modeling Project as product owner \u2014 delivered the first DevEx predictive models (viscosity, HMW), producing 7 dashboards, 2 production models, and ~2 FTE savings while impacting 6 major portfolio programs.')
bullet('Directed C3.AI document extraction pilot \u2014 structured data from >3,000 documents with >95% accuracy \u2014 then guided transition to in-house platform for long-term scalability aligned with IT strategy.')
bullet('Championed SPARTAN (Domain-Driven Design platform) adoption across PD, recognized with an individual STELA Award. Deployed ROVER as the first SPARTAN production tool in PD.')

sub_heading('BRIDGING IT & DRUG DEVELOPMENT')
bullet('Served as DPD-IT Enterprise Team Lead \u2014 established governance for software prioritization, vendor negotiation, and budget management. Achieved full PD funding for business-critical software.')
bullet('Negotiated IT resources for PD DevEx projects, establishing patterns the broader DevEx Modeling team subsequently adopted. Drove alignment between IT architecture standards and scientific speed-to-value needs.')
bullet('Commissioned the DPD Data Map and developed the Data Maturity Framework \u2014 informing the PD Data Strategy and enabling visibility into digital gaps across the organization.')

sub_heading('PORTFOLIO GOVERNANCE & PROGRAM MANAGEMENT')
bullet('Sponsored and governed ~20 digital/automation initiatives \u2014 ensuring business-case rigor, milestone tracking, resource alignment, stakeholder engagement, and sustained adoption through change management.')
bullet('As DevEx Modeling Core Team Lead, initiated 4 new modeling projects (FEM, DEM, PIV SDD nozzle, CM RTD), revitalized Biopharm Digital II, and implemented biweekly value realization meetings \u2014 capturing 190+ experiments and 27+ assets leveraging modeling.')
bullet('Pivoted eBatch record initiative when vendor failed \u2014 led transition to in-house "Tabula" platform, resulting in $300K YoY cost savings while maintaining program timelines.')

sub_heading('INNOVATION & EMERGING TOOLS')
bullet('Led DPD Innovation Champion Team: scaled pipeline from 3 to 15 projects (2023\u21922024), drove 9 Innovation Council funding decisions, and established ROI-based evaluation frameworks.')
bullet('Launched DPD DevEx Spotlight Series to build visibility around AI, modeling, and platform innovations \u2014 aligning with enterprise goals of speed to FIH and resource efficiency.')
bullet('Expanded DevLIMS as structured data capture tool across biologics, powder testing, and MSE workflows \u2014 automating data entry for 90%+ of analytical techniques.')

sub_heading('TEAM DEVELOPMENT & LEADERSHIP')
bullet('Lead a multidisciplinary organization of 7 direct reports (including 3 PhDs) plus 3 people managers across automation, data science, PAT, and modeling.')
bullet('Developed Manager Readiness framework and role-competency models; grew 3 direct reports into new managers in a single year. MyVoice engagement score 83 vs. BMS average 75 (+8), Inclusion +11.')

# --- ROLE 2: Associate Director / Principal Scientist ---
role_header('Associate Director / Principal Scientist, Digital Strategy & Automation', 'Mar 2019\u2013Jun 2022')
org_line('Engineering Technologies, Drug Product Science & Technology')
role_summary(
    'Created Digital Engineering as a new discipline within DPST, defining roles, career paths, and a cohesive strategy. '
    'Key member of the ET leadership team driving organizational strategy and culture.'
)

bullet('Co-led the GPS Digital Moonshot initiative \u2014 spearheaded a PD/CTD effort engaging 100 leaders, scientists, and partners to produce a 5-year digital vision with 21+ minimum viable products identified for investment.')
bullet('Developed three-pillar digital strategy: (1) scalable data platforms with UX focus, (2) multi-modal information access, (3) next-gen development and manufacturing technologies.')
bullet('Directed a digital portfolio of cross-modality projects including Lumetics, CIFER, MAST, ARGILE AR/VR (28 projects, 17 sites \u2014 instrumental during pandemic lockdown), and Nisaba (14+ instruments).')
bullet('Initiated DPD-IT Enterprise Strategy team, establishing the governance and partnership model between DPD and IT that persists today.')
bullet('As DPD culture team leader and BOLD member, drove D&I initiatives and served as panelist for PD leadership forums.')

# --- ROLE 3: SRS I -> II ---
role_header('Senior Research Scientist I \u2192 II, Automation & Digitalization', '2013\u20132019')
org_line('Engineering Technologies / Drug Product Science & Technology  (SRS I: 2013\u20132017 | SRS II: 2017\u20132019)')
role_summary(
    'First formal technical leadership scope: managed 6 scientists in the automation function, coordinated the department '
    'automation matrix team, and began driving digital strategy alongside automation expansion.'
)

bullet('Envisioned, built the business case for, and led a multimillion-dollar automation function that grew from 1 capability to 13 workflows serving ~25 studies/year and impacting 80+ assets. Delivered ~5.05 FTE total savings.')
bullet('Led $3M biologics automation program (vendor selection through deployment) \u2014 biologics formulation screening at scale (180+ combinations, 1,000+ samples, 10,000+ data points).')
bullet('Developed the Three-Tier Automation Value Framework: (1) Productivity savings, (2) Data & predictive insight, (3) New scientist-technology experimentation paradigms \u2014 guiding all subsequent investment decisions.')
bullet('Deployed automated buffer-exchange workflow (1.1 FTE savings) that influenced vendor Freeslate (now Unchained Labs) to invest in a commercial platform based on BMS-developed technology.')
bullet('Recognized externally through AAPS webinars, user-group presentations, and co-authored publications on lab automation strategy.')

# --- ROLE 4: Preformulation ---
role_header('Preformulation Scientist & Automation Developer', '2003\u20132013')
org_line('Exploratory Pharmaceutics / Drug Product Science & Technology')

bullet('Led drug product teams across modalities \u2014 OSD, IV, biologics \u2014 including Fostemsavir (HIV), Beclabuvir (HCV), Milvexian (Factor XIa), supporting IND/IND-enabling milestones.')
bullet('Built foundational automation from scratch: Automated Excipient Compatibility (32 projects, 2.7 FTE savings), viscous liquid solubility screening (30 projects, 1.25 FTE savings; 75% faster, ~4\u00d7 coverage).')
bullet('Developed mathematical models for polymorphic phase impurity impact on bioavailability; expanded dissolution/absorption simulations using Matlab.')

# =========================================================================
# EDUCATION
# =========================================================================
section_heading('EDUCATION')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('B.S.E., Chemical Engineering')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = dark
run.font.name = 'Calibri'
run2 = p.add_run(' \u2014 Princeton University, Princeton, NJ (2003)')
run2.font.size = Pt(10)
run2.font.color.rgb = medium
run2.font.name = 'Calibri'

# =========================================================================
# PUBLICATIONS
# =========================================================================
section_heading('PUBLICATIONS')

def pub_entry(text):
    """Add a numbered publication entry."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = dark
    run.font.name = 'Calibri'

pub_entry('1.  Ying, W.; Levons, J.; Carney, A.; Gandhi, R.; Vydra, V.; Rubin, E. \u201cSemi-Automated Sample Preparation for Protein Stability and Formulation Screening via Buffer Exchange.\u201d JALA (2015), 21(3), 378.')
pub_entry('2.  Hemenway, J.; Carvalho, T.; Rao, V.; Wu, Y.; Levons, J.; Narang, A.; Paruchuri, S.; Stamato, H.; Varia, S. \u201cFormation of Reactive Impurities in Aqueous and Neat Polyethylene Glycol 400 and Effects of Antioxidants and Oxidation Inducers.\u201d Journal of Pharmaceutical Sciences (2012), 101(9), 3305\u20133318.')
pub_entry('3.  Wu, Y.; Levons, J.; Narang, A.; Raghavan, K.; Rao, V. \u201cReactive Impurities in Excipients: Profiling, Identification and Mitigation of Drug\u2013Excipient Incompatibility.\u201d AAPS PharmSciTech (2011), 12(4), 1248\u20131263.')
pub_entry('4.  Chandran, S.; Gesenberg, C.; Levons, J.; Hubert, M.; Raghavan, K. \u201cA High-Throughput Spectrophotometric Approach for Evaluation of Precipitation Resistance.\u201d Journal of Pharmaceutical and Biomedical Analysis (2011), 56(4), 698\u2013704.')
pub_entry('5.  Li, J.; Yang, B.; Levons, J.; Pinnamaneni, S.; Raghavan, K. \u201cPhase Behavior of TPGS\u2013PEG400/1450 Systems and Their Application to Liquid Formulation: A Formulation Platform Approach.\u201d Journal of Pharmaceutical Sciences (2011), 100(11), 4907\u20134921.')
pub_entry('6.  Gu, C.; Li, H.; Levons, J.; Lentz, K.; Gandhi, R.; Raghavan, K.; Smith, R. \u201cPredicting Effect of Food on Extent of Drug Absorption Based on Physicochemical Properties.\u201d Pharmaceutical Research (2007), 24(6), 1118\u20131130.')

# =========================================================================
# PATENTS
# =========================================================================
section_heading('PATENTS')

pub_entry('1.  Gu, C-H.; Gao, Q.; Kuang, S-M.; Lai, C.; Levons, J.K.; Qian, F. \u201cFormulations of 1-(4-benzoylpiperazin-1-yl)-2-[4-methoxy-7-(3-methyl-[1,2,4]triazol-1-yl)-1H-pyrrolo[2,3-c]pyridin-3-yl]ethane-1,2-dione.\u201d U.S. Pat. Appl. Publ. US 2006100209 A1 (2006).')
pub_entry('2.  Chen, C-P.H.; Digiugno, D.; Gao, Q.; Gu, C-H.; Levons, J.K.; Yang, B-S. \u201cCrystalline Forms of 1-benzoyl-4-[2-[4-methoxy-7-(3-methyl-1H-1,2,4-triazol-1-yl)-1-[(phosphonooxy)methyl]-1H-pyrrolo[2,3-c]pyridin-3-yl]-1,2-dioxoethyl]piperazine.\u201d PCT Int. Appl. WO 2007070589 A2 (2007).')
pub_entry('3.  Chandran, S.; Gandhi, R.B.; Levons, J.K.; Perrone, R.K.; Price, C.P.; Raghavan, K.S.; Ullah, I. \u201cBioavailable Capsule Compositions of Amorphous \u03b1-(N-sulfonamido)acetamide Compound for Treatment of CNS Disorders.\u201d U.S. Pat. Appl. Publ. US 20100260837 A1 (2010).')

# =========================================================================
# RECOGNITION
# =========================================================================
section_heading('RECOGNITION & PROFESSIONAL ENGAGEMENT')
bullet('STELA Award (Individual, 2024) \u2014 SPARTAN / Domain-Driven Design adoption across PD')
bullet('AAPS webinar presenter and automation user-group speaker')
bullet('Vendor co-development with Symyx/Freeslate \u2192 commercially deployed platform industry-wide')
bullet('Active BOLD member; D&I advocate and panelist across BMS/PD forums')
bullet('Manager ratings: \u201cPhenomenal performance\u201d (2024), \u201cTop performance\u201d (2025)')

# =========================================================================
# SAVE
# =========================================================================
output_path = r'C:\Users\jaqua\Downloads\JKL_Resumes\Levons_BMS_ExecDir_AI_Enablement_Resume.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
