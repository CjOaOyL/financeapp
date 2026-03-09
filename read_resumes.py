import os, glob

try:
    from docx import Document
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document

folder = r'C:\Users\jaqua\Downloads\JKL_Resumes\JKL_Resumes\Recent Resumes'
output_file = r'C:\Users\jaqua\Calculus\all_resumes_text.txt'
docs = sorted(glob.glob(os.path.join(folder, '*.docx')))
sep = '=' * 80

with open(output_file, 'w', encoding='utf-8') as out:
    for doc_path in docs:
        fname = os.path.basename(doc_path)
        out.write(f'\n{sep}\n')
        out.write(f'FILE: {fname}\n')
        out.write(f'{sep}\n')
        try:
            doc = Document(doc_path)
            for para in doc.paragraphs:
                out.write(para.text + '\n')
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    out.write(' | '.join(cells) + '\n')
        except Exception as e:
            out.write(f'ERROR reading {fname}: {e}\n')

print(f'Done. Written to {output_file}')
size = os.path.getsize(output_file)
print(f'File size: {size:,} bytes')
